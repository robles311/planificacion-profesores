#!/usr/bin/env python3
"""
generate_docx.py — Genera el Word (.docx) de planificación semanal para Fran.

Uso:
    python3 tools/generate_docx.py --input output/content.json --output output/planificacion.docx

Input: JSON con estructura de planificación (ver schema abajo).
Output: archivo .docx con el formato exacto de la plantilla de Fran.
"""

import argparse
import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ────────────────────────────────────────────────────────────
# Helpers de formato
# ────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex: str):
    """Pinta el fondo de una celda con color hexadecimal (sin #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def cell_text(cell, text: str, bold=False, size=10, color=None, align=None):
    """Escribe texto en una celda, limpiando contenido previo."""
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))


def add_header_row(table, label: str, value: str, label_color='2E4057', value_color=None):
    """Agrega una fila de cabecera label | value a una tabla de 2 columnas."""
    row = table.add_row()
    lc = row.cells[0]
    vc = row.cells[1]
    set_cell_bg(lc, label_color)
    cell_text(lc, label, bold=True, size=10, color='FFFFFF')
    cell_text(vc, value, bold=False, size=10)


def set_table_borders(table):
    """Agrega bordes simples a toda la tabla."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'AAAAAA')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_section_heading(doc, text: str, level=2, color='2E4057'):
    """Agrega un párrafo de encabezado de sección con color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11 if level == 1 else 10)
    run.font.color.rgb = RGBColor(*bytes.fromhex(color))


def add_field(doc, label: str, value: str):
    """Agrega un campo label: valor en el documento."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(f'{label}: ')
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_value = p.add_run(value)
    run_value.font.size = Pt(10)


def add_bullet_list(doc, items: list):
    """Agrega una lista de bullets."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(item)
        run.font.size = Pt(10)


def add_language_levels(doc, fonologico: str, semantico: str, pragmatico: str):
    """Agrega la tabla de niveles de lenguaje al final de cada ficha."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('NIVELES DE LENGUAJE')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)  # azul oscuro

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    set_table_borders(table)

    # Encabezados
    headers = ['NIVEL FONOLÓGICO', 'NIVEL SEMÁNTICO', 'NIVEL PRAGMÁTICO']
    header_colors = ['2874A6', '1A5276', '154360']
    hdr_row = table.rows[0]
    for i, (h, c) in enumerate(zip(headers, header_colors)):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, c)
        cell_text(cell, h, bold=True, size=9, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER)

    # Contenido
    content_row = table.add_row()
    for i, content in enumerate([fonologico, semantico, pragmatico]):
        cell = content_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(content)
        run.font.size = Pt(9)

    doc.add_paragraph()


def add_page_break(doc):
    """Agrega un salto de página."""
    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Construcción del documento
# ────────────────────────────────────────────────────────────

def build_cabecera(doc, cabecera: dict):
    """Tabla de cabecera con datos del curso."""
    add_section_heading(doc, 'PLANIFICACIÓN SEMANAL', level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    set_table_borders(table)

    # Ancho de columnas
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)

    campos = [
        ('Mes', cabecera.get('mes', '')),
        ('Nivel', cabecera.get('nivel', '')),
        ('Educadora', cabecera.get('educadora', '')),
        ('Temática', cabecera.get('tematica', '')),
        ('Semana', cabecera.get('semana', '')),
    ]
    for label, value in campos:
        row = table.add_row()
        lc, vc = row.cells[0], row.cells[1]
        set_cell_bg(lc, '2E4057')
        cell_text(lc, label, bold=True, size=10, color='FFFFFF')
        cell_text(vc, value, size=10)

    doc.add_paragraph()


def build_estaciones_llegada(doc, estaciones_data: dict):
    """Sección de estaciones de llegada (08:45–09:15)."""
    if not estaciones_data:
        return

    add_section_heading(doc, 'ESTACIONES DE LLEGADA (08:45 a 09:15)', level=1)

    desc = estaciones_data.get('descripcion', 'Juego simbólico en estaciones')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(desc)
    run.font.size = Pt(10)
    run.italic = True

    estaciones = estaciones_data.get('estaciones', [])
    if not estaciones:
        doc.add_paragraph()
        return

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    set_table_borders(table)

    hdr = table.rows[0]
    set_cell_bg(hdr.cells[0], '2E4057')
    cell_text(hdr.cells[0], 'Estación', bold=True, size=9, color='FFFFFF')
    set_cell_bg(hdr.cells[1], '2E4057')
    cell_text(hdr.cells[1], 'Material / Actividad', bold=True, size=9, color='FFFFFF')

    for est in estaciones:
        row = table.add_row()
        cell_text(row.cells[0], est.get('nombre', ''), bold=True, size=9)
        cell_text(row.cells[1], est.get('material', ''), size=9)

    nota = estaciones_data.get('nota', '')
    if nota:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        r = p.add_run(f'Nota: {nota}')
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()


def build_resumen_semanal(doc, resumen: dict):
    """Tabla de resumen semanal (5 días × bloques)."""
    add_section_heading(doc, 'RESUMEN SEMANAL', level=1)

    dias_orden = ['lunes', 'martes', 'miercoles', 'jueves', 'viernes']
    dias_labels = {
        'lunes': 'LUNES',
        'martes': 'MARTES',
        'miercoles': 'MIÉRCOLES',
        'jueves': 'JUEVES',
        'viernes': 'VIERNES',
    }

    dias_presentes = [d for d in dias_orden if d in resumen]
    if not dias_presentes:
        return

    # Calcular max bloques
    max_bloques = max(len(resumen[d].get('bloques', [])) for d in dias_presentes)

    # Tabla: 1 fila header + max_bloques filas
    n_cols = len(dias_presentes)
    table = doc.add_table(rows=1 + max_bloques, cols=n_cols)
    table.style = 'Table Grid'
    set_table_borders(table)

    # Encabezados con emoji
    for col_i, dia in enumerate(dias_presentes):
        cell = table.rows[0].cells[col_i]
        emoji = resumen[dia].get('emoji', '')
        fecha = resumen[dia].get('fecha', dias_labels.get(dia, dia))
        set_cell_bg(cell, '2E4057')
        cell_text(cell, f'{emoji} {fecha}', bold=True, size=9, color='FFFFFF',
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # Contenido de bloques
    for col_i, dia in enumerate(dias_presentes):
        bloques = resumen[dia].get('bloques', [])
        for b_i, bloque in enumerate(bloques):
            if 1 + b_i < len(table.rows):
                cell = table.rows[1 + b_i].cells[col_i]
                tipo = bloque.get('tipo', '')
                tema = bloque.get('tema', '')
                horario = bloque.get('horario', '')
                texto = f'{tipo}\n{tema}\n{horario}' if tema else f'{tipo}\n{horario}'
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(tipo)
                run.bold = True
                run.font.size = Pt(9)
                if tema:
                    p.add_run(f'\n{tema}').font.size = Pt(9)
                if horario:
                    hr = cell.add_paragraph()
                    r = hr.add_run(horario)
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()


def build_bloque(doc, dia_label: str, bloque: dict, primer_bloque_del_dia: bool):
    """Crea la ficha detallada de un bloque."""
    if primer_bloque_del_dia:
        add_section_heading(doc, f'DÍA: {dia_label}', level=1, color='1A3A4A')
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run('—' * 80)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    num = bloque.get('numero', '')
    horario = bloque.get('horario', '')
    tipo = bloque.get('tipo', '')

    # Encabezado del bloque
    heading_text = f'Bloque {num} — {horario} · {tipo}'
    add_section_heading(doc, heading_text, level=2, color='2874A6')

    # Meta
    add_field(doc, 'Meta', bloque.get('meta', ''))

    # Tabla OA / OAT / Habilidad
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    set_table_borders(table)
    hdr = table.rows[0]
    for i, label in enumerate(['Ámbito / Núcleo', 'OA', 'OAT', 'Habilidad']):
        set_cell_bg(hdr.cells[i], '2E4057')
        cell_text(hdr.cells[i], label, bold=True, size=9, color='FFFFFF')

    row = table.add_row()
    cell_text(row.cells[0], bloque.get('ambito_nucleo', ''), size=9)
    cell_text(row.cells[1], bloque.get('oa', ''), size=9)
    cell_text(row.cells[2], bloque.get('oat', ''), size=9)
    cell_text(row.cells[3], bloque.get('habilidad', ''), size=9)
    doc.add_paragraph()

    # Actividad
    add_section_heading(doc, 'Actividad', level=2, color='2E4057')
    inicio = bloque.get('inicio', '')
    desarrollo = bloque.get('desarrollo', '')
    cierre = bloque.get('cierre', '')

    if inicio:
        p = doc.add_paragraph()
        r = p.add_run('Inicio: ')
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(inicio).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    if desarrollo:
        p = doc.add_paragraph()
        r = p.add_run('Desarrollo: ')
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(desarrollo).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    if cierre:
        p = doc.add_paragraph()
        r = p.add_run('Cierre: ')
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(cierre).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    # Indicadores de logro
    add_section_heading(doc, 'Indicadores de logros', level=2, color='2E4057')
    indicadores = bloque.get('indicadores', [])
    add_bullet_list(doc, indicadores)

    # Evaluación
    eval_proc = bloque.get('evaluacion_procedimiento', 'Observación')
    eval_inst = bloque.get('evaluacion_instrumento', 'Lista de apreciación')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run('Evaluación — ')
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(f'Procedimiento: {eval_proc} · Instrumento: {eval_inst}').font.size = Pt(10)

    # Recursos
    add_section_heading(doc, 'Recursos', level=2, color='2E4057')
    recursos = bloque.get('recursos', [])
    add_bullet_list(doc, recursos)

    # Niveles de lenguaje
    add_language_levels(
        doc,
        bloque.get('nivel_fonologico', ''),
        bloque.get('nivel_semantico', ''),
        bloque.get('nivel_pragmatico', ''),
    )

    # Separador entre bloques
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


def generate(content: dict, output_path: str):
    """Genera el .docx completo a partir del JSON de contenido."""
    doc = Document()

    # Márgenes del documento
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Estilo base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # 1. Cabecera
    build_cabecera(doc, content.get('cabecera', {}))

    # 2. Estaciones de llegada (si están definidas)
    if content.get('estaciones_llegada'):
        build_estaciones_llegada(doc, content['estaciones_llegada'])

    # 3. Resumen semanal
    build_resumen_semanal(doc, content.get('resumen', {}))

    # 3. Fichas por día y bloque
    dias = content.get('dias', [])
    for dia_data in dias:
        dia_label = dia_data.get('dia', '')
        bloques = dia_data.get('bloques', [])
        for i, bloque in enumerate(bloques):
            build_bloque(doc, dia_label, bloque, primer_bloque_del_dia=(i == 0))

    doc.save(output_path)
    print(f'[OK] Documento generado: {output_path}')


# ────────────────────────────────────────────────────────────
# Entry point
# ────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Genera el .docx de planificación semanal')
    parser.add_argument('--input', required=True, help='Ruta al JSON de contenido')
    parser.add_argument('--output', required=True, help='Ruta de salida del .docx')
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f'[ERROR] No se encontró el archivo de entrada: {input_path}', file=sys.stderr)
        sys.exit(1)

    with open(input_path, encoding='utf-8') as f:
        content = json.load(f)

    Path(args.output).parent.mkdir(parents=True, exist_ok=True)
    generate(content, args.output)


if __name__ == '__main__':
    main()
